import streamlit as st
from fpdf import FPDF
from docx import Document
import pypandoc
import pandas as pd
import re
from zipfile import ZipFile
from io import BytesIO
import os
import shutil
from pdf2docx import Converter

# Función para cargar la hoja de estilos CSS
def load_css(file_name):
    with open(file_name) as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

# Función para leer el documento de Word y buscar palabras entre {{ }}
def leer_documento_y_buscar_palabras(doc_path):
    doc = Document(doc_path)
    palabras = set()
    pattern = re.compile(r"\{\{(.*?)\}\}")

    for para in doc.paragraphs:
        matches = pattern.findall(para.text)
        for match in matches:
            palabras.add(match.strip())

    return list(palabras)

# Función para reemplazar palabras en el documento de Word
def reemplazar_palabras(doc_path, output_path, reemplazos):
    doc = Document(doc_path)
    pattern = re.compile(r"\{\{(.*?)\}\}")

    for para in doc.paragraphs:
        para_text = para.text
        for match in pattern.findall(para_text):
            clave = match.strip()
            if clave in reemplazos:
                para_text = para_text.replace(f"{{{{{match}}}}}", str(reemplazos[clave]))
        para.text = para_text

    doc.save(output_path)

# Función para convertir docx a pdf usando pypandoc con xelatex
def convertir_docx_a_pdf(docx_path, pdf_path, progress_bar=None):
    try:
        pypandoc.convert_file(docx_path, 'pdf', outputfile=pdf_path, extra_args=['--pdf-engine=xelatex'])
    except OSError:
        pypandoc.download_pandoc()
        pypandoc.convert_file(docx_path, 'pdf', outputfile=pdf_path, extra_args=['--pdf-engine=xelatex'])
    if progress_bar:
        progress_bar.progress(100)

# Función para convertir pdf a docx usando pdf2docx
def convertir_pdf_a_docx(pdf_path, docx_path, progress_bar=None):
    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)
    cv.close()
    if progress_bar:
        progress_bar.progress(100)

# Función para leer los valores desde el archivo Excel o CSV
def leer_valores_desde_archivo(file_path):
    if file_path.endswith('.xlsx'):
        df = pd.read_excel(file_path)
    elif file_path.endswith('.csv'):
        df = pd.read_csv(file_path)
    else:
        raise ValueError("El archivo debe ser un .xlsx o .csv")
    return df

# Función para guardar archivos subidos
def guardar_archivo_subido(uploaded_file, save_path):
    with open(save_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

# Función para cargar archivos del histórico
def cargar_historico(nombre_historico):
    doc_path = os.path.join(nombre_historico, "documento.docx")
    archivo_path = None
    for ext in ['xlsx', 'csv']:
        path = os.path.join(nombre_historico, f"valores.{ext}")
        if os.path.exists(path):
            archivo_path = path
            break
    return doc_path, archivo_path

# Función para eliminar un histórico
def eliminar_historico(nombre_historico):
    shutil.rmtree(nombre_historico)

# Función principal de Streamlit
def main():
    load_css("styles.css")
    st.title("Generador de PDF desde Formulario Dinámico")

    menu = ["Generar PDF desde Formulario", "Convertir DOCX a PDF", "Convertir PDF a DOCX"]
    choice = st.sidebar.selectbox("Selecciona una opción", menu)

    if choice == "Generar PDF desde Formulario":
        st.markdown("""
            **Proceso:**
            1. Sube un archivo de Word con marcadores de posición en el formato `{{nombre}}`.
            2. Sube un archivo de Excel o CSV con las columnas correspondientes a los marcadores de posición.
            3. Genera documentos PDF con los valores sustituidos y descárgalos.
        """)

        # Listar los historiales de exportaciones disponibles excluyendo la carpeta 'ejemplos'
        historiales = [d for d in os.listdir() if os.path.isdir(d) and not d.startswith('.') and d != 'ejemplos']
        historial_seleccionado = st.selectbox("Selecciona un registro del histórico de exportaciones", ["Nuevo"] + historiales)

        if historial_seleccionado != "Nuevo":
            if st.button("Eliminar histórico"):
                eliminar_historico(historial_seleccionado)
                st.success(f"Histórico {historial_seleccionado} eliminado con éxito.")
                st.experimental_rerun()

            doc_path, archivo_path = cargar_historico(historial_seleccionado)

            if archivo_path:
                palabras = leer_documento_y_buscar_palabras(doc_path)
                df_valores = leer_valores_desde_archivo(archivo_path)

                if palabras and set(palabras).issubset(df_valores.columns):
                    st.write(f"Histórico seleccionado: {historial_seleccionado}")

                    fuente = st.selectbox("Selecciona la fuente del texto", ["Arial", "Courier", "Helvetica", "Times"])
                    tamano = st.slider("Selecciona el tamaño del texto", 8, 24, 12)
                    nombre_documento = st.text_input("Introduce el nombre base para los documentos PDF", historial_seleccionado)
                    columna_nombre = st.selectbox("Selecciona la columna para el nombre del archivo PDF", df_valores.columns)

                    if st.button("Descargar PDFs existentes"):
                        zip_buffer = BytesIO()
                        with ZipFile(zip_buffer, 'a') as zf:
                            for file in os.listdir(historial_seleccionado):
                                if file.endswith('.pdf'):
                                    with open(os.path.join(historial_seleccionado, file), "rb") as f:
                                        zf.writestr(file, f.read())

                        st.download_button(
                            label="Descargar PDFs",
                            data=zip_buffer.getvalue(),
                            file_name=f"{historial_seleccionado}.zip",
                            mime="application/zip"
                        )

                    if st.button("Generar nuevos PDFs con nuevas opciones"):
                        zip_buffer = BytesIO()
                        with ZipFile(zip_buffer, 'a') as zf:
                            progress_bar = st.progress(0)
                            total_files = len(df_valores)
                            for index, row in df_valores.iterrows():
                                form_data = {palabra: str(row[palabra]) for palabra in palabras}
                                
                                # Crear un nuevo documento de Word con los valores reemplazados
                                reemplazo_doc_path = os.path.join(historial_seleccionado, f"{nombre_documento}_{index+1}.docx")
                                reemplazar_palabras(doc_path, reemplazo_doc_path, form_data)
                                
                                # Convertir el documento de Word a PDF
                                reemplazo_pdf_path = os.path.join(historial_seleccionado, f"{nombre_documento}-{str(row[columna_nombre])}.pdf")
                                convertir_docx_a_pdf(reemplazo_doc_path, reemplazo_pdf_path)
                                
                                with open(reemplazo_pdf_path, "rb") as pdf_file:
                                    zf.writestr(f"{nombre_documento}-{str(row[columna_nombre])}.pdf", pdf_file.read())
                                
                                progress_bar.progress(int((index + 1) / total_files * 100))

                        st.success("Nuevos PDFs generados con éxito")

                        st.download_button(
                            label="Descargar PDFs",
                            data=zip_buffer.getvalue(),
                            file_name=f"{historial_seleccionado}.zip",
                            mime="application/zip"
                        )

                    # Opción para subir archivos y sobrescribir los existentes
                    st.write("Sube nuevos archivos para sobrescribir los existentes:")

                    uploaded_file_doc = st.file_uploader("Sube un nuevo archivo de Word", type=["docx"], key="new_docx")
                    uploaded_file_archivo = st.file_uploader("Sube un nuevo archivo de Excel o CSV con los valores", type=["xlsx", "csv"], key="new_archivo")

                    if uploaded_file_doc:
                        guardar_archivo_subido(uploaded_file_doc, doc_path)
                        st.success("Archivo de Word sobrescrito con éxito.")

                    if uploaded_file_archivo:
                        guardar_archivo_subido(uploaded_file_archivo, archivo_path)
                        st.success("Archivo de Excel/CSV sobrescrito con éxito.")
                else:
                    st.warning("Las columnas en el archivo Excel/CSV no coinciden con las palabras clave encontradas en el documento de Word o no se encontraron palabras entre {{ }} en el documento.")
            else:
                st.warning("No se encontró un archivo Excel o CSV en el histórico.")
        else:
            uploaded_file_doc = st.file_uploader("Sube un archivo de Word", type=["docx"])
            uploaded_file_archivo = st.file_uploader("Sube un archivo de Excel o CSV con los valores", type=["xlsx", "csv"])
            nombre_historico = st.text_input("Introduce el nombre para el histórico de exportaciones")

            fuentes = ["Arial", "Courier", "Helvetica", "Times"]
            tamano_default = 12

            if uploaded_file_doc and uploaded_file_archivo and nombre_historico:
                # Crear directorio para el histórico de exportaciones
                os.makedirs(nombre_historico, exist_ok=True)

                # Guardar archivos subidos
                doc_path = os.path.join(nombre_historico, "documento.docx")
                archivo_path = os.path.join(nombre_historico, f"valores.{uploaded_file_archivo.name.split('.')[-1]}")
                guardar_archivo_subido(uploaded_file_doc, doc_path)
                guardar_archivo_subido(uploaded_file_archivo, archivo_path)

                palabras = leer_documento_y_buscar_palabras(doc_path)
                df_valores = leer_valores_desde_archivo(archivo_path)

                if palabras:
                    st.write("Por favor, revisa y ajusta los valores si es necesario:")

                    if set(palabras).issubset(df_valores.columns):
                        fuente = st.selectbox("Selecciona la fuente del texto", fuentes)
                        tamano = st.slider("Selecciona el tamaño del texto", 8, 24, tamano_default)
                        nombre_documento = st.text_input("Introduce el nombre base para los documentos PDF", nombre_historico)
                        columna_nombre = st.selectbox("Selecciona la columna para el nombre del archivo PDF", df_valores.columns)

                        with st.form("formulario"):
                            form_data_list = []
                            for index, row in df_valores.iterrows():
                                with st.expander(f"Registro {index+1}"):
                                    form_data = {}
                                    for palabra in palabras:
                                        form_data[palabra] = st.text_input(f"{palabra}", value=str(row[palabra]))
                                    form_data_list.append(form_data)
                            submit = st.form_submit_button("Generar PDFs")

                        if submit:
                            zip_buffer = BytesIO()
                            with ZipFile(zip_buffer, 'a') as zf:
                                progress_bar = st.progress(0)
                                total_files = len(form_data_list)
                                for index, form_data in enumerate(form_data_list):
                                    # Crear un nuevo documento de Word con los valores reemplazados
                                    reemplazo_doc_path = os.path.join(nombre_historico, f"{nombre_documento}_{index+1}.docx")
                                    reemplazar_palabras(doc_path, reemplazo_doc_path, form_data)
                                    
                                    # Convertir el documento de Word a PDF
                                    reemplazo_pdf_path = os.path.join(nombre_historico, f"{nombre_documento}-{str(df_valores.at[index, columna_nombre])}.pdf")
                                    convertir_docx_a_pdf(reemplazo_doc_path, reemplazo_pdf_path)
                                    
                                    with open(reemplazo_pdf_path, "rb") as pdf_file:
                                        zf.writestr(f"{nombre_documento}-{str(df_valores.at[index, columna_nombre])}.pdf", pdf_file.read())
                                    
                                    progress_bar.progress(int((index + 1) / total_files * 100))

                            st.success("PDFs generados con éxito")

                            st.download_button(
                                label="Descargar PDFs",
                                data=zip_buffer.getvalue(),
                                file_name=f"{nombre_historico}.zip",
                                mime="application/zip"
                            )
                    else:
                        st.warning("Las columnas en el archivo Excel/CSV no coinciden con las palabras clave encontradas en el documento de Word.")
                else:
                    st.warning("No se encontraron palabras entre {{ }} en el documento.")
            elif not nombre_historico:
                st.info("Por favor, introduce un nombre para el histórico de exportaciones.")
            else:
                st.info("Por favor, sube un archivo de Word y un archivo de Excel/CSV para continuar.")

        # Proporcionar ejemplos de archivos en un ZIP
        st.header("Descargar ejemplos")

        if st.button("Descargar ejemplos en un ZIP"):
            ejemplo_dir = "ejemplos"
            zip_buffer = BytesIO()
            with ZipFile(zip_buffer, 'w') as zf:
                ejemplo_files = ["ejemplo.docx", "ejemplo.xlsx", "ejemplo.csv"]
                for file_name in ejemplo_files:
                    file_path = os.path.join(ejemplo_dir, file_name)
                    zf.write(file_path, arcname=file_name)
            
            st.download_button(
                label="Descargar ejemplos",
                data=zip_buffer.getvalue(),
                file_name="ejemplos.zip",
                mime="application/zip"
            )

    elif choice == "Convertir DOCX a PDF":
        st.header("Convertir DOCX a PDF")
        uploaded_file_docx = st.file_uploader("Sube un archivo DOCX", type=["docx"])
        if uploaded_file_docx:
            if not os.path.exists("uploads"):
                os.makedirs("uploads")
            docx_path = os.path.join("uploads", uploaded_file_docx.name)
            guardar_archivo_subido(uploaded_file_docx, docx_path)
            pdf_path = docx_path.replace(".docx", ".pdf")
            progress_bar = st.progress(0)
            convertir_docx_a_pdf(docx_path, pdf_path, progress_bar)
            with open(pdf_path, "rb") as f:
                st.download_button(
                    label="Descargar PDF",
                    data=f,
                    file_name=os.path.basename(pdf_path),
                    mime="application/pdf"
                )
            os.remove(docx_path)
            os.remove(pdf_path)

    elif choice == "Convertir PDF a DOCX":
        st.header("Convertir PDF a DOCX")
        uploaded_file_pdf = st.file_uploader("Sube un archivo PDF", type=["pdf"])
        if uploaded_file_pdf:
            if not os.path.exists("uploads"):
                os.makedirs("uploads")
            pdf_path = os.path.join("uploads", uploaded_file_pdf.name)
            guardar_archivo_subido(uploaded_file_pdf, pdf_path)
            docx_path = pdf_path.replace(".pdf", ".docx")
            progress_bar = st.progress(0)
            convertir_pdf_a_docx(pdf_path, docx_path, progress_bar)
            with open(docx_path, "rb") as f:
                st.download_button(
                    label="Descargar DOCX",
                    data=f,
                    file_name=os.path.basename(docx_path),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            os.remove(pdf_path)
            os.remove(docx_path)

if __name__ == "__main__":
    main()
